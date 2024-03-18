from typing import Iterator
from docx import Document as DocxDocument
from langchain.docstore.document import Document
from langchain.document_loaders.base import BaseBlobParser
from langchain.document_loaders.blob_loaders.schema import Blob



    
class DOCXParser(BaseBlobParser):
    def lazy_parse(self, blob: Blob) -> Iterator[Document]:
        with blob.as_bytes_io() as file:
            doc = DocxDocument(file)
            content = ""
            for paragraph in doc.paragraphs:
                content += f"{paragraph.text}\n"
            metadata = {}
            yield Document(page_content=content, metadata=metadata)
